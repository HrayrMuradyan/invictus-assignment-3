import sys
import os
sys.path.append(os.path.join(os.path.dirname(__file__), ".."))
from src.config import StyleConfig
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

class CoverPageProcessor:
    def __init__(self, doc):
        self.doc = doc

    def process(self):
        """
        Processes the cover page of the document.
        """
        self._normalize_vertical_alignment()
        self._apply_text_styling()

    def _is_row_blank(self, item):
        """
        Returns True if the paragraph or XML element contains no visible text.
        """
        if item is None:
            return True
    
        text = ""
        
        # It's a paragraph text
        if hasattr(item, 'text'):
            text = item.text
            
        # It's a low-level XML Element
        elif hasattr(item, 'findall'):
            t_tags = item.findall('.//w:t', namespaces=item.nsmap)
            text = "".join([t.text for t in t_tags if t.text])
            
        return not bool(text and text.strip())

    def _set_font(self, paragraph, size):
        """
        Set the font (provided in the StyleConfig) and size of a paragraph.
        """
        for run in paragraph.runs:
            run.font.name = StyleConfig.FONT_NAME
            run.font.size = size

    def _normalize_vertical_alignment(self):
        """
        Ensures the first row with text is row 19.
        """
        all_paragraphs = self.doc.paragraphs

        # Find the index of the first non-empty paragraphfirst_test_index
        first_text_index = 0
        found_text = False
        
        for i, p in enumerate(all_paragraphs):
            if not self._is_row_blank(p):
                first_text_index = i
                found_text = True
                break

        if not found_text:
            first_text_index = len(all_paragraphs)

        # Calculate how many blank rows have to be inserted
        required_index = StyleConfig.COVER_START_ROW - 1

        if first_text_index < required_index:
            missing_lines = required_index - first_text_index
            for _ in range(missing_lines):
                p = self.doc.paragraphs[0].insert_paragraph_before("")
                # Ensure blank rows are the required size
                self._set_font(p, StyleConfig.FONT_SIZE)

        # If the number of whitespaces is larger than the required index
        elif first_text_index > required_index:
            # We need to remove some lines
            lines_to_remove = first_text_index - required_index
            
            # We remove from the top 
            removed_count = 0
            
            # Use a while loop because the paragraphs list changes size as we delete (for loop would break)
            while removed_count < lines_to_remove:
                # Always look at the current first paragraph
                # Break if the doc becomes empty
                if not self.doc.paragraphs: 
                    break 
                
                p_to_remove = self.doc.paragraphs[0]
                
                # Only delete if it is actually empty
                if self._is_row_blank(p_to_remove):
                    p_element = p_to_remove._element
                    p_element.getparent().remove(p_element)
                    p_to_remove._p = p_to_remove._element = None
                    
                    removed_count += 1
                else:
                    # We hit text before we finished deleting. Stop to prevent data loss.
                    # Although it's unlikely, it's a good safety test
                    print(f"Warning: Stopped deleting rows. Found unexpected text at row {removed_count}.")
                    break

    def _format_company_title(self, paragraph):
        """
        Format the company title of the cover page (First line)
        """
        full_text = paragraph.text.strip()
        paragraph.clear()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Splits the title to four groups (defined below)
        match = re.search(r'^(.*?)(\(formerly)(.*?)(\))$', full_text, re.IGNORECASE)
        
        if match:
            parts = [
                (match.group(1), True),  # Company Name -> CAPS
                (match.group(2), False), # (formerly -> lower
                (match.group(3), True),  # Old Name -> Title Case
                (match.group(4), False)  # )
            ]
        else:
            # if regex fails, it will just capitalize the title
            # this is just a safety check as the instructions define
            # the correct title structure, which has to match the regex above
            parts = [(full_text, True)]

        for text_part, is_name in parts:
            run = paragraph.add_run(text_part)
            run.font.name = StyleConfig.FONT_NAME
            run.font.size = StyleConfig.COVER_TITLE_SIZE
            run.font.bold = True

            # Formerly
            if "formerly" in text_part.lower():
                run.text = text_part.lower()

            # Company name
            elif is_name:
                # If main company name, it's CAPS
                if text_part == match.group(1):
                    run.text = text_part.upper()
                else:
                    run.text = text_part.title()

    def _enforce_one_blank_row_after(self, current_paragraph):
        """
        This function ensures that after each text line there is one blank row
        """
        next_element = current_paragraph._element.getnext()

        # End of document
        if next_element is None:
            new_p = self.doc.add_paragraph("") 
            self._set_font(new_p, StyleConfig.FONT_SIZE)
            return

        
        if next_element.tag.endswith('p'):
            # If next element is text
            if not self._is_row_blank(next_element):
                # Insert a blank row before this text paragraph (after our current one)
                next_p_obj = self.doc.paragraphs[self.doc.element.body.index(next_element)]
                new_p = next_p_obj.insert_paragraph_before("")
                self._set_font(new_p, StyleConfig.FONT_SIZE)
                
            # if next line is a blank row
            else:
                sibling_of_next = next_element.getnext()

                while sibling_of_next is not None and \
                      sibling_of_next.tag.endswith('p') and \
                      self._is_row_blank(sibling_of_next):

                    # Remove the second blank row
                    parent = sibling_of_next.getparent()
                    parent.remove(sibling_of_next)

                    # Update the pointer to check the next one
                    sibling_of_next = next_element.getnext()

    def _apply_text_styling(self):
        """
        Adjusts the text styling for the cover page including the title and other components.
        """
        i = 0
        n_rows_to_check = 30
        while i < n_rows_to_check and i < len(self.doc.paragraphs):
            p = self.doc.paragraphs[i]
            text = p.text.strip()

            date_regex = r'\b(January|February|March|April|May|June|July|August|September|October|November|December)\b.*\d{4}'

            # Title
            # Look for "... (formerly ...)" pattern
            if re.search(r"\(formerly\b", text, re.IGNORECASE):
                self._format_company_title(p)
                self._enforce_one_blank_row_after(p)
                
                # Skip the next blank row
                i+=1
                
            # Second line
            elif "financial statements" in text.lower():
                p.style = self.doc.styles['Normal']
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.name = StyleConfig.FONT_NAME
                    run.font.size = StyleConfig.FONT_SIZE
                    run.font.bold = True
                    # Capitalize Each Word
                    run.text = run.text.title() 
                    
                self._enforce_one_blank_row_after(p)
                i+=1
                
            # Third line (period)
            elif re.search(date_regex, text, re.IGNORECASE):
                p.style = self.doc.styles['Normal']
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.bold = True
                    run.font.name = StyleConfig.FONT_NAME
                    run.font.size = StyleConfig.FONT_SIZE
                    
                self._enforce_one_blank_row_after(p)
                i+=1

            # Fourth line
            elif "unaudited" in text.lower() or "expressed in" in text.lower():
                p.style = self.doc.styles['Normal']
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                result = text[0] + text[1].upper() + text[2:]

                p.text = result
            
                for run in p.runs:
                    run.font.name = StyleConfig.FONT_NAME
                    run.font.size = StyleConfig.FONT_SIZE
                    run.font.bold = False
            
            i+=1