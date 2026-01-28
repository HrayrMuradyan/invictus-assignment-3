from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re
import sys
import os
sys.path.append(os.path.join(os.path.dirname(__file__), ".."))
from src.config import StyleConfig
from docx.enum.table import WD_ROW_HEIGHT_RULE
import logging
logger = logging.getLogger(__name__)

class TableProcessor:
    def __init__(self, doc):
        self.doc = doc

    def _set_cell_margins(self, table):
        """
        Directly modifies table XML to set cell margins to 0.05cm (L/R) and 0cm (T/B).
        """
        tbl_pr = table._tbl.tblPr
        
        # Create the table cell margin element
        tbl_cell_mar = OxmlElement('w:tblCellMar')
        
        for side in ['left', 'right']:
            node = OxmlElement(f'w:{side}')
            node.set(qn('w:w'), "28") 
            node.set(qn('w:type'), 'dxa')
            tbl_cell_mar.append(node)
            
        for side in ['top', 'bottom']:
            node = OxmlElement(f'w:{side}')
            node.set(qn('w:w'), "0")
            node.set(qn('w:type'), 'dxa')
            tbl_cell_mar.append(node)
    
        # Remove old margins if they exist and append new ones
        if tbl_pr.find(qn('w:tblCellMar')) is not None:
            tbl_pr.remove(tbl_pr.find(qn('w:tblCellMar')))
        tbl_pr.append(tbl_cell_mar)

    def process(self):
        """
        Process each table in the table page
        """
        for table in self.doc.tables:
            self._apply_structural_rules(table)
            self._apply_semantic_bolding(table)

    def _identify_current_period_column(self, table):
        """
        Given two "Value" columns, identify the current period column
        """

        # Maps col_index -> datetime object
        dates = {}

        for _, row in enumerate(table.rows[:5]):
            for col_idx, cell in enumerate(row.cells):
                text = cell.text.strip()

                # Find years like 2023, 2024... 2099
                match = re.search(r'\b(20\d{2})\b', text)
                
                if match:
                    try:
                        year = int(match.group(1))
                        dates[col_idx] = year
                    except ValueError:
                        pass

        if not dates:
            logger.warning("Couldn't identify the current period column.")
            return None

        return max(dates, key=dates.get)

    def _apply_semantic_bolding(self, table):
        """
        Bolds the 'Current Period' value column and UN-BOLDS the 'Prior Period' value column.
        """
        current_period_col_idx = self._identify_current_period_column(table)
        
        if current_period_col_idx is None:
            logger.warning("Because no current period column was identified, the bolding wasn't applied")
            return

        for row in table.rows:
            # Skip malformed rows (e.g. merged headers)
            if current_period_col_idx >= len(row.cells):
                continue
                
            for col_idx, cell in enumerate(row.cells):
                # Strictly ignore columns 0 and 1.
                if col_idx < 2:
                    continue

                # If this IS the current period column -> Force BOLD
                # If this IS NOT the current period column -> Force UN-BOLD
                should_be_bold = (col_idx == current_period_col_idx)
                
                text = cell.text.strip()
                
                # Apply only if the cell has the following content -> (numbers, $)
                has_data = any(c.isdigit() for c in text) or '$' in text
                
                if has_data:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.font.bold = should_be_bold

    def _force_table_xml_grid(self, table):
        """
        Manually rewrites the <w:tblGrid> XML element.
        Word uses this grid to strictly enforce column widths.
        """
        # Access the internal XML element for the table
        tbl = table._tbl
        
        # Find or Create the 'tblGrid' element
        tblGrid = tbl.tblGrid
        if tblGrid is None:
            tblGrid = OxmlElement('w:tblGrid')
            tbl.insert(0, tblGrid)
        else:
            # Clear existing grid columns to prevent conflicts
            tblGrid.clear()
            
        # Build new grid columns based on Config
        for width_emu in StyleConfig.TABLE_COLUMN_WIDTHS:
            gridCol = OxmlElement('w:gridCol')
            
            # Convert EMU (python-docx default) to Twips (XML default)
            width_twips = str(int(width_emu / 635))
            
            gridCol.set(qn('w:w'), width_twips)
            tblGrid.append(gridCol)
                
    def _apply_structural_rules(self, table):
        """
        Apply the structural rules to the table, including cell margins, column widths, font and identation
        """
        # Force the Table Grid XML
        self._force_table_xml_grid(table)

        self._set_cell_margins(table)

        for row in table.rows:
            row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            row.height = StyleConfig.TABLE_ROW_HEIGHT

            for idx, cell in enumerate(row.cells):
                # We still set cell.width for redundancy, though the Grid usually wins
                if idx < len(StyleConfig.TABLE_COLUMN_WIDTHS):
                    cell.width = StyleConfig.TABLE_COLUMN_WIDTHS[idx]
                
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = StyleConfig.FONT_NAME
                        run.font.size = StyleConfig.FONT_SIZE

                    # Hanging Indent Logic
                    if idx == 0 and len(cell.text.strip()) > 0:
                        paragraph.paragraph_format.left_indent = StyleConfig.TABLE_HANGING_INDENT
                        paragraph.paragraph_format.first_line_indent = -StyleConfig.TABLE_HANGING_INDENT
                        

            
            