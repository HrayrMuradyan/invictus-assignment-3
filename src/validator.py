import logging
from docx import Document
from docx.document import Document as DocumentClass
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

# Configure logging to show distinct PASS/FAIL status
logging.basicConfig(level=logging.INFO, format='%(message)s')
logger = logging.getLogger("Validator")

def validate_output(doc_or_path):
    """
    Validates a processed .docx file against specific style guide rules.
    """

    issues = []

    if isinstance(doc_or_path, (str, Path)):
        try:
            doc = Document(str(doc_or_path))
        except Exception as e:
            message = f"Could not load document. {e}"
            logger.error(message)
            issues.append(message)
            return issues

    elif isinstance(doc_or_path, DocumentClass):
        doc = doc_or_path

    else:
        message = (
            "The argument 'doc_or_path' must be either a path to a .docx file "
            f"or a python-docx Document object. Got {type(doc_or_path)}"
        )
        logger.error(message)
        issues.append(message)
        return issues

    logger.info("Starting Validation\n")
    
    # Helper to find the next non-empty paragraph (text)
    def get_next_content_node(start_index, paragraphs):
        for i in range(start_index, len(paragraphs)):
            if paragraphs[i].text.strip():
                return i, paragraphs[i]
        return -1, None

    # --- Section 1: Cover Page ---
    logger.info("Section 1: Cover Page\n")
    
    # Find the first line (Title)
    title_index, title_p = get_next_content_node(0, doc.paragraphs)
    
    # Check: Start Row (Row 19) [Rule: 17]
    if title_index >= 0 and title_index == 18:
        logger.info("[PASS] Cover page text starts on Row 19.")
    else:
        message = f"[FAIL] Cover page text starts on Row {title_index + 1} (Expected: 19)."
        logger.warning(message)
        issues.append(message)

    # Check: Company Title Style [Rules: 20, 21]
    # Requirements: Bold, Size 14, Centered
    if title_p:
        is_bold = all(run.font.bold for run in title_p.runs if run.text.strip())
        is_size_14 = any(run.font.size == Pt(14) for run in title_p.runs if run.font.size)
        is_centered = title_p.alignment == WD_ALIGN_PARAGRAPH.CENTER
        
        if is_bold and is_size_14 and is_centered:
            logger.info("[PASS] LINE 1 (Title): Correctly Bold, Size 14, and Centered.")
        else:
            message = (
                f"[FAIL] LINE 1 (Title): Style mismatch. Bold: {is_bold}, "
                f"Size 14: {is_size_14}, Centered: {is_centered}"
            )
            logger.warning(message)
            issues.append(message)

        # Check 3: Blank Row After Title [Rule: 25]
        if title_index + 1 < len(doc.paragraphs) and not doc.paragraphs[title_index + 1].text.strip():
            logger.info("[PASS] Blank row exists after Title.")
        else:
            message = "[FAIL] Missing blank row after Title."
            logger.warning(message)
            issues.append(message)
    else:
        message = "CRITICAL: No text found on cover page."
        logger.error(message)
        issues.append(message)
        return issues

    # Find Second Line (Financial Statements)
    # Start searching after the title + blank row
    stmts_index, stmts_p = get_next_content_node(title_index + 1, doc.paragraphs)

    # Check: Financial Statements Style [Rule: 26]
    # Requirements: Bold, Capitalize Each Word, Centered
    if stmts_p:
        text = stmts_p.text.strip()
        is_bold = any(run.font.bold for run in stmts_p.runs)
        is_title_case = text == text.title()
        
        if "financial statements" in text.lower():
            if is_bold and is_title_case:
                logger.info("[PASS] LINE 2: '%s' is Bold and Title Case.", text)
            else:
                message = f"[FAIL] LINE 2: '{text}' style mismatch. Bold: {is_bold}, Title Case: {is_title_case}"
                logger.warning(message)
                issues.append(message)
        else:
            message = f"[FAIL] LINE 2: Expected 'Financial Statements', found '{text}'." 
            logger.warning(message)
            issues.append(message)

        # Check: Blank Row After [Rule: 27]
        if stmts_index + 1 < len(doc.paragraphs) and not doc.paragraphs[stmts_index + 1].text.strip():
             logger.info("[PASS] Blank row exists after Financial Statements.")
        else:
             message = "[FAIL] Missing blank row after Financial Statements."
             logger.warning(message)
             issues.append(message)
    else:
        message = "[FAIL] LINE 2 missing."
        logger.warning(message)
        issues.append(message)

    # Find Third Line (Period Reference)
    period_index, period_p = get_next_content_node(stmts_index + 1, doc.paragraphs)

    # Check: Period Reference Style [Rule: 28]
    # Requirements: Bold, Sentence case (only first letter cap)
    if period_p:
        text = period_p.text.strip()
        is_bold = any(run.font.bold for run in period_p.runs)
        
        # Check if not CAPS
        is_not_all_caps = not text.isupper() 

        if is_bold and is_not_all_caps:
             logger.info("[PASS] LINE 3: '%s' is Bold.", text)
        else:
             message = f"[FAIL] LINE 3: '{text}' style mismatch. Bold: {is_bold}, Caps Check: {is_not_all_caps}"
             logger.warning(message)
             issues.append(message)

        # Check: Blank Row After [Rule: 29]
        if period_index + 1 < len(doc.paragraphs) and not doc.paragraphs[period_index + 1].text.strip():
             logger.info("[PASS] Blank row exists after Period Reference.")
        else:
             message = "[FAIL] Missing blank row after Period Reference."
             logger.warning(message)
             issues.append(message)
    else:
        message = "[FAIL] LINE 3 missing."
        logger.warning(message)
        issues.append(message)

    # Find Fourth Line (Unaudited)
    _, unaudited_p = get_next_content_node(period_index + 1, doc.paragraphs)

    # Check: Unaudited Style [Rule: 30]
    # Requirements: Sentence case, NOT Bold (implied standard font), Centered
    if unaudited_p:
        text = unaudited_p.text.strip()
        
        # Check for "(Unaudited...)"
        if ("unaudited" in text.lower()) or ("expressed" in text.lower()):
            # Check Bold (Should be False)
            is_bold = any(run.font.bold for run in unaudited_p.runs)
            
            # Needs clarification here
            # Flexible check: ensure it's not ALL CAPS or Title Case
            clean_text = text.strip()
            is_sentence_ish = (
                clean_text.startswith("(")
                and len(clean_text) > 2
                and clean_text[1].isupper()
                and clean_text[2].islower()
                and not clean_text.istitle() 
            )

            if not is_bold and is_sentence_ish:
                logger.info("[PASS] LINE 4: '%s' is Un-bolded and Sentence Case.\n", text)
            else:
                message = f"[FAIL] LINE 4: '{text}' style mismatch. Bold: {is_bold} (Should be False).\n"
                logger.warning(message)
                issues.append(message)
        else:
            message = f"[FAIL] LINE 4: Expected 'Unaudited...' or 'Expressed...', found '{text}'.\n" 
            logger.warning(message)
            issues.append(message)
    else:
        message = "[FAIL] LINE 4 missing.\n"
        logger.warning(message)
        issues.append(message)

    # --- Section 2: Table Validation ---
    logger.info("Section 2: Tables\n")
    
    if not doc.tables:
        message = "No tables found in document."
        logger.warning(message)
        issues.append(message)

    for t_idx, table in enumerate(doc.tables):
        logger.info("Checking Table %s...", t_idx + 1)

        # Check: Row Height (Rule: At least 0.37cm) [Rule: 32]
        rows_pass = True
        for row in table.rows:
            if row.height is None or row.height.cm < 0.37:
                rows_pass = False
                break
        
        if rows_pass:
            logger.info("   [PASS] Row heights are at least 0.37cm.")
        else:
            message = "   [FAIL] One or more rows have incorrect height."
            logger.warning(message)
            issues.append(message)

        # Check: Cell Margins (Rule: L/R 0.05cm, T/B 0.0cm) 
        tbl_pr = table._tbl.tblPr
        mar = tbl_pr.find(qn('w:tblCellMar'))
        if mar is not None:
            left = mar.find(qn('w:left'))
            top = mar.find(qn('w:top'))
            
            # 28 dxa is approx 0.05cm
            if left is not None and left.get(qn('w:w')) == '28' and top is not None and top.get(qn('w:w')) == '0':
                 logger.info("   [PASS] Cell margins are set correctly (0.05cm L/R, 0cm T/B).")
            else:
                 message = "   [FAIL] Cell margins in XML do not match expected values."
                 logger.warning(message)
                 issues.append(message)
        else:
            message = "   [FAIL] No custom cell margins found in XML."
            logger.warning(message)
            issues.append(message)

        # Check: Column Widths [Rules: 34, 37]
        # Expected: ~11.99cm, ~1.20cm, ~2.30cm (converted to Twips or EMUs)
        # 1 cm = 360000 EMU. 
        expected_widths = [11.99, 1.20, 2.30, 2.30]
        
        # Check the first row's cells
        width_pass = True
        first_row_cells = table.rows[0].cells
        for i, expected in enumerate(expected_widths):
            if i < len(first_row_cells):
                # Allow small tolerance for floating point conversion
                cell_width_cm = first_row_cells[i].width.cm
                if abs(cell_width_cm - expected) > 0.1: 
                    width_pass = False
                    message = f"   [FAIL] Col {i} width mismatch. Found {round(cell_width_cm, 2)}cm, Expected {expected}cm."
                    logger.warning(message)
                    issues.append(message)
        
        if width_pass:
            logger.info("   [PASS] Column widths match specifications.")

        # Check: Hanging Indent (Rule: 0.63cm) [Rule: 35]
        hanging_passed = False
        checked_row = -1
        
        for r_idx, row in enumerate(table.rows):
            # Skip header (row 0)
            if r_idx == 0:
                continue
                
            cell = row.cells[0]
            # Must have text to have indent
            if cell.text.strip(): 
                p = cell.paragraphs[0]
                fmt = p.paragraph_format
                
                # Check if values exist
                left_val = fmt.left_indent.cm if fmt.left_indent else 0.0
                first_line_val = fmt.first_line_indent.cm if fmt.first_line_indent else 0.0
                
                # Verify 0.63cm hanging (Left +0.63, First Line -0.63)
                if (abs(left_val - 0.63) < 0.05 and abs(first_line_val + 0.63) < 0.05):
                    hanging_passed = True
                    checked_row = r_idx
                    break 
                else:
                    # Capture the failure details for the log
                    checked_row = r_idx
                    break 

        if hanging_passed:
             logger.info(f"   [PASS] Hanging indent detected in Row {checked_row+1}, Col 1.\n")
        elif checked_row != -1:
             # We found a text row, but it had wrong indentation
             p = table.rows[checked_row].cells[0].paragraphs[0]
             fmt = p.paragraph_format
             l_val = fmt.left_indent.cm if fmt.left_indent else "None"
             fl_val = fmt.first_line_indent.cm if fmt.first_line_indent else "None"
             
             message = (
                 f"   [FAIL] Hanging indent mismatch in Row {checked_row+1}. "
                 f"Found Left={l_val}, FirstLine={fl_val}. Expected Left=0.63, FirstLine=-0.63.\n"
             )
             logger.warning(message)
             issues.append(message)
        else:
             # We never found a row with text
             logger.warning("   [SKIP] Could not validate hanging indent (No data rows found in Col 1).\n")

    # --- Global Font Check ---
    logger.info("Section 3: Global Font\n")
    
    font_issues_found = False

    for i, p in enumerate(doc.paragraphs):

        # Skip the Title (Arial 14)
        if i == title_index:
            continue
            
        # Skip blank rows
        if not p.text.strip():
            continue

        # Check every 'run' in the paragraph
        for run in p.runs:
            # Skip runs without text
            if not run.text.strip():
                continue

            # Check Font Name: Must be "Arial"
            name_bad = (run.font.name != "Arial")
            
            # Check Font Size: Must be 9pt (if explicitly set)
            size_bad = (run.font.size is not None and run.font.size != Pt(9))
            if name_bad or size_bad:
                font_issues_found = True
                
                actual_name = run.font.name if run.font.name else "None (Inherited)"
                actual_size = run.font.size.pt if run.font.size else "None (Inherited)"
                
                # Truncate text for cleaner logs
                preview = (p.text[:30] + '...') if len(p.text) > 30 else p.text
                
                message = (
                    f"   [FAIL] Font Issue in Para {i+1} ('{preview}'): "
                    f"Name='{actual_name}', Size={actual_size}. Expected Arial 9pt."
                )
                logger.warning(message)
                issues.append(message)
                
                # Stop checking runs in this paragraph to avoid spam
                break

    # Check table
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                for p in cell.paragraphs:
                    
                    if not p.text.strip():
                        continue

                    for run in p.runs:
                        if not run.text.strip():
                            continue
                        name_bad = (run.font.name != "Arial")
                        size_bad = (run.font.size is not None and run.font.size != Pt(9))

                        if name_bad or size_bad:
                            print("HIMA PROBLEM CHKA?")
                            font_issues_found = True
                            actual_name = run.font.name if run.font.name else "None (Inherited)"
                            actual_size = run.font.size.pt if run.font.size else "None (Inherited)"
                            preview = (p.text[:20] + '...') if len(p.text) > 20 else p.text
                            
                            message = (
                                f"   [FAIL] Table Font Issue (Table {t_idx+1}, Row {r_idx+1}, Col {c_idx+1}): "
                                f"'{preview}' -> Name='{actual_name}', Size={actual_size}."
                            )
                            logger.warning(message)
                            issues.append(message)
                            
                            break

    if not font_issues_found:
        logger.info("[PASS] Body font appears to be Arial 9pt.\n")
    else:
        message = "[FAIL] Body font is not Arial 9pt.\n"
        logger.warning(message)
        issues.append(message)

    logger.info("VALIDATION COMPLETE\n")

    return issues